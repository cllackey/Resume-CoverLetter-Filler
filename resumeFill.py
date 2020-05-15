from tkinter import *
from tkinter import filedialog
from docx import Document
import pickle
import os
import sys
import comtypes.client
import ctypes, sys
from datetime import date


def is_admin():
    try:
        return ctypes.windll.shell32.IsUserAnAdmin()
    except:
        return False



# if is_admin():
root = Tk()
root.title = "Resume Filler"
root.wm_title("Resume Filler")
root.geometry("620x410+300+150")
res_open = ""
cov_open = ""
gpa_val = ""
entry_width = 80
addresses = ["", "", ""]
copy = ""
copy_val = 0
try:
    addresses = pickle.load(open("save.p", "rb" ))
    print(addresses[0])
    res_open = addresses[0]
    cov_open = addresses[1]
    gpa_val = addresses[2]
except (OSError, IOError) as e:
    res_open = ""
    cov_open = ""
except IndexError as e:
    try:
        res_open = addresses[0]
    except IndexError as e1:
        res_open = ""
    try:
        cov_open = addresses[1]
    except IndexError as e1:
        cov_open = ""
    try:
        gpa_val = addresses[2]
    except IndexError as e1:
        gpa_val = ""
    addresses = [res_open, cov_open, gpa_val]
    pickle.dump(addresses, open("save.p", "wb"))


def make_copy(filepath, pop):
    global copy_val
    global copy
    while True:
        copy_val = copy_val + 1
        copy = "(" + str(copy_val) + ")"
        if os.path.exists(filepath + "/" + comp_name.get() + "/" + "caden_lackey_resume_" +
                     comp_name.get().replace(" ", "") + copy + ".docx"):
            continue
        if os.path.exists(filepath + "/" + comp_name.get() + "/" + "caden_lackey_cover_" +
                     comp_name.get().replace(" ", "") + copy + ".docx"):
            continue
        break
    pop.quit()
    pop.destroy()


def replace_file(filepath, pop):
    try:
        os.remove(filepath + "/" + comp_name.get() + "/" + "caden_lackey_resume_" +
                  comp_name.get().replace(" ", "") + ".docx")
        os.remove(filepath + "/" + comp_name.get() + "/" + "caden_lackey_cover_" +
                  comp_name.get().replace(" ", "") + ".docx")
    except PermissionError:
        pop_error = Tk()
        pop_error.title = "!"
        pop_error.wm_title("!")
        pop_error.geometry("200x70+510+305")
        pop_error_label = Label(pop_error, text="Permission Error\nTry running as Administrator")
        pop_error_label.pack()
        pop_error_butt = Button(pop_error, text="Okay", command=pop_error.destroy)
        pop_error_butt.pack()
        pop_error.mainloop()


def res_select():
    res_open = filedialog.askopenfilename(initialdir="Documents", title="Select file",
                                          filetypes=(("Word files", "*.docx"), ("all files", "*.*")))
    print("res0: " + res_open)
    addresses[0] = res_open
    print("addr[0]: " + addresses[0])
    pickle.dump(addresses, open("save.p", "wb"))
    res_loc_text.delete(0, END)
    res_loc_text.insert(0, res_open)


def cov_select():
    cov_open = filedialog.askopenfilename(initialdir="/", title="Select file",
                                          filetypes=(("Word files", "*.docx"), ("all files", "*.*")))
    addresses[1] = cov_open
    pickle.dump(addresses, open("save.p", "wb"))
    cov_loc_text.delete(0, END)
    cov_loc_text.insert(0, cov_open)


def create_files():
    print("res: " + res_open)
    pop1 = Tk()
    pop1.title = "!"
    pop1.wm_title("!")
    pop1.geometry("200x70+510+305")
    pop1_label = Label(pop1, text="Creating Files. Please Wait.")
    pop1_label.pack()
    while True:
        pop1.update_idletasks()
        pop1.update()
        addresses[2] = gpa.get()
        addresses[1] = cov_open
        addresses[0] = res_open
        pickle.dump(addresses, open("save.p", "wb"))
        res_doc = Document(res_open)
        print("AFter Opem")
        # os.remove(res_open)
        for paragraph in res_doc.paragraphs:
            if '[POSITION]' in paragraph.text:
                paragraph.text = paragraph.text.replace("[POSITION]", tkvar.get() + " " + job_pos.get())
            if '[COMPANY]' in paragraph.text:
                paragraph.text = paragraph.text.replace("[COMPANY]", comp_name.get())
        cov_doc = Document(cov_open)
        # os.remove(cov_open)
        for paragraph in cov_doc.paragraphs:
            if '[DATE]' in paragraph.text:
                paragraph.text = paragraph.text.replace("[DATE]", date_text.get())
            if '[HIRINGMANAGER]' in paragraph.text:
                paragraph.text = paragraph.text.replace("[HIRINGMANAGER]", hire_man.get())
            if '[ADDRESS1]' in paragraph.text:
                paragraph.text = paragraph.text.replace("[ADDRESS1]", addr_1.get())
            if '[ADDRESS2]' in paragraph.text:
                paragraph.text = paragraph.text.replace("[ADDRESS2]", addr_2.get())
            if '[POSITION]' in paragraph.text:
                paragraph.text = paragraph.text.replace("[POSITION]", job_pos.get() + " position")
            if '[COMPANY]' in paragraph.text:
                paragraph.text = paragraph.text.replace("[COMPANY]", comp_name.get())
            if '[GPA]' in paragraph.text:
                paragraph.text = paragraph.text.replace("[GPA]", gpa.get())
        i = len(res_open) - 1
        while i >= 0:
            if res_open[i] == "/":
                break
            i = i - 1
        filepath = res_open[0:i]
        try:
            os.mkdir(filepath + "/" + comp_name.get())
        except OSError:
            print("Path Not Created.")
        else:
            print("Path Created")
        while True:
            try:
                res_doc.save(filepath + "/" + comp_name.get() + "/" + "caden_lackey_resume_" +
                             comp_name.get().replace(" ", "") + copy + ".docx")
                cov_doc.save(filepath + "/" + comp_name.get() + "/" + "caden_lackey_cover_" +
                             comp_name.get().replace(" ", "") + copy + ".docx")
                # try:
                #     pop1.destroy()
                # except:
                #     print("NOthing to destry")
                break
            except PermissionError as e0:
                pop_exists = Tk()
                pop_exists.title = "!"
                pop_exists.wm_title("!")
                pop_exists.geometry("200x70+510+305")
                pop_exists_label = Label(pop_exists, text="File Already Exists")
                pop_exists_label.pack()
                # pop_exists_butt0 = Button(pop_exists, text="Replace",
                #                           command=lambda: replace_file(filepath, pop_exists))
                # pop_exists_butt0.grid(row=1, column=0)
                pop_exists_butt1 = Button(pop_exists, text="Create New",
                                          command=lambda: make_copy(filepath, pop_exists))
                pop_exists_butt1.pack()
                pop_exists.mainloop()
                # pop_error = Tk()
                # pop_error.title = "!"
                # pop_error.wm_title("!")
                # pop_error.geometry("200x70+510+305")
                # pop_error_label = Label(pop_error, text="Permission Error\nTry running as Administrator")
                # pop_error_label.pack()
                # pop_error_butt = Button(pop_error, text="Okay", command=pop_error.destroy)
                # pop_error_butt.pack()
                # pop_error.mainloop()
        print("Reached")
        wdFormatPDF = 17
        in_file1 = filepath + "/" + comp_name.get() + "/" + "caden_lackey_resume_" + \
                   comp_name.get().replace(" ", "") + copy + ".docx"
        out_file1 = filepath + "/" + comp_name.get() + "/" + "caden_lackey_resume_" + \
                    comp_name.get().replace(" ", "") + copy + ".pdf"
        in_file2 = filepath + "/" + comp_name.get() + "/" + "caden_lackey_cover_" + \
                   comp_name.get().replace(" ", "") + copy + ".docx"
        out_file2 = filepath + "/" + comp_name.get() + "/" + "caden_lackey_cover_" + \
                    comp_name.get().replace(" ", "") + copy + ".pdf"
        word = comtypes.client.CreateObject('Word.Application')
        word.Visible = False
        res_doc2 = word.Documents.Open(in_file1)
        res_doc2.SaveAs(out_file1, FileFormat=wdFormatPDF)
        res_doc2.Close()
        cov_doc2 = word.Documents.Open(in_file2)
        cov_doc2.SaveAs(out_file2, FileFormat=wdFormatPDF)
        cov_doc2.Close()
        word.Quit()
        pop1.destroy()
        break
    pop2 = Tk()
    pop2.title = "Complete"
    pop2.wm_title("Complete")
    pop2.geometry("200x70+510+305")
    pop2_label = Label(pop2, text="Files Created!")
    pop2_label.pack()
    pop2_butt = Button(pop2, text="Okay", command=pop2.destroy)
    pop2_butt.pack()
    pop2.mainloop()


res_loc_text = Entry(root, width=entry_width)
res_loc_text.grid(row=0, column=1, columnspan=2, padx=10, pady=10)
cov_loc_text = Entry(root, width=entry_width)
cov_loc_text.grid(row=1, column=1, columnspan=2, padx=10, pady=10)

res_loc_text.delete(0, END)
res_loc_text.insert(0, addresses[0])
cov_loc_text.delete(0, END)
cov_loc_text.insert(0, addresses[1])

res_butt = Button(root, text="Select File", command=res_select)
res_butt.grid(row=0, column=0)
cov_butt = Button(root, text="Select File", command=cov_select)
cov_butt.grid(row=1, column=0)

get_gpa = Label(root, text="GPA:")
get_gpa.grid(row=2, column=0, padx=10, pady=10)
gpa = Entry(root, width=entry_width)
gpa.grid(row=2, column=1, columnspan=2, padx=10, pady=10)
gpa.delete(0, END)
gpa.insert(0, gpa_val)

get_comp_name = Label(root, text="Company Name:")
get_comp_name.grid(row=3, column=0, padx=10, pady=10)
comp_name = Entry(root, width=entry_width)
comp_name.grid(row=3, column=1, columnspan=2, padx=10, pady=10)

get_job_pos = Label(root, text="Job Position:")
get_job_pos.grid(row=4, column=0, padx=10, pady=10)
job_pos = Entry(root, width=70)
job_pos.grid(row=4, column=2, columnspan=1, pady=10)

tkvar = StringVar(root)
choices = {'a', 'an'}
tkvar.set('a')
pos_menu = OptionMenu(root, tkvar, *choices)
pos_menu.grid(row=4, column=1)

get_date = Label(root, text="Date:")
get_date.grid(row=5, column=0, padx=10, pady=10)
date_text = Entry(root, width=entry_width)
date_text.grid(row=5, column=1, columnspan=2, padx=10, pady=10)
today = date.today()
article = "th"
if today.strftime("%d") == "1":
    article = "st"
elif today.strftime("%d") == "2":
    article = "nd"
elif today.strftime("%d") == "3":
    article = "rd"
curr_date = today.strftime("%B %d" + article + ", %Y")
date_text.delete(0, END)
date_text.insert(0, curr_date)

get_hire_man = Label(root, text="Hiring Manger:")
get_hire_man.grid(row=6, column=0, padx=10, pady=10)
hire_man = Entry(root, width=entry_width)
hire_man.grid(row=6, column=1, columnspan=2, padx=10, pady=10)

get_addr_1 = Label(root, text="Address Line 1:")
get_addr_1.grid(row=7, column=0, padx=10, pady=10)
addr_1 = Entry(root, width=entry_width)
addr_1.grid(row=7, column=1, columnspan=2, padx=10, pady=10)

get_addr_2 = Label(root, text="Address Line 1:")
get_addr_2.grid(row=8, column=0, padx=10, pady=10)
addr_2 = Entry(root, width=entry_width)
addr_2.grid(row=8, column=1, columnspan=2, padx=10, pady=10)

create_butt = Button(root, text="Create Files", command=create_files)
create_butt.grid(row=10, column=0, columnspan=3, pady=10)

while True:
    root.update_idletasks()
    root.update()
# else:
#     ctypes.windll.shell32.ShellExecuteW(None, "runas", sys.executable, " ".join(sys.argv), None, 1)